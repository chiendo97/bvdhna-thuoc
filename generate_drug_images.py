#!/usr/bin/env python3
"""Generate responsive HTML files for each drug from DOCX source."""

import re
from html import escape
from pathlib import Path

from docx import Document
from docx.table import Table, _Cell

BASE_DIR = Path(__file__).parent
DOCX_FILE = BASE_DIR / "data.docx"
OUTPUT_DIR = BASE_DIR / "docs" / "drugs"

# Column headers for the table
COLUMN_HEADERS = [
    "Tên hoạt chất",
    "Dược thư Quốc gia Việt Nam 2022",
    "Sanford guide (update 12.2025)",
    "HD hiệu chỉnh liều ở BN suy thận - BV Bạch Mai 2023",
    "Renal Pharmacotherapy 2021",
]


def sanitize_filename(name: str) -> str:
    """Convert drug name to a safe filename."""
    # Replace non-breaking space and regular space
    name = name.replace('\xa0', ' ')
    name = re.sub(r'[<>:"/\\|?*]', '_', name)
    name = name.replace(' ', '_')
    name = re.sub(r'_+', '_', name)
    return name.strip('_')


def get_cell_span(cell: _Cell) -> tuple[int, int]:
    """Get the column span and row span for a cell."""
    tc = cell._tc
    # Get grid span (colspan)
    grid_span = tc.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}gridSpan")
    if grid_span is not None:
        colspan = int(grid_span)
    else:
        # Check tcPr for gridSpan
        tc_pr = tc.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcPr")
        if tc_pr is not None:
            grid_span_el = tc_pr.find(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}gridSpan"
            )
            if grid_span_el is not None:
                colspan = int(
                    grid_span_el.get(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "1"
                    )
                )
            else:
                colspan = 1
        else:
            colspan = 1

    # Row span is more complex - for now just return 1
    rowspan = 1
    return colspan, rowspan


def extract_nested_table_html(table: Table) -> str:
    """Convert a nested DOCX table to HTML, handling merged cells."""
    html_parts = ['<table class="nested-table">']

    for row in table.rows:
        html_parts.append("<tr>")
        seen_cells = set()  # Track cells we've already processed (by their XML element id)

        for cell in row.cells:
            # Skip if we've already processed this cell (merged cells appear multiple times)
            cell_id = id(cell._tc)
            if cell_id in seen_cells:
                continue
            seen_cells.add(cell_id)

            # Get colspan
            colspan, rowspan = get_cell_span(cell)

            cell_text = escape(cell.text.strip())
            # Handle line breaks within cells
            cell_text = cell_text.replace("\n", "<br>")

            # Build td with colspan if needed
            if colspan > 1:
                html_parts.append(f'<td colspan="{colspan}">{cell_text}</td>')
            else:
                html_parts.append(f"<td>{cell_text}</td>")

        html_parts.append("</tr>")

    html_parts.append("</table>")
    return "".join(html_parts)


# XML namespaces for Office Math
MATH_NS = {'m': 'http://schemas.openxmlformats.org/officeDocument/2006/math',
           'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}


def extract_math_html(math_element) -> str:
    """Convert Office Math (OMML) element to HTML representation."""
    # Check for fractions
    fracs = math_element.findall('.//m:f', namespaces=MATH_NS)
    if fracs:
        result_parts = []
        for frac in fracs:
            num_texts = frac.findall('.//m:num//m:t', namespaces=MATH_NS)
            den_texts = frac.findall('.//m:den//m:t', namespaces=MATH_NS)
            num = ''.join([t.text or '' for t in num_texts])
            den = ''.join([t.text or '' for t in den_texts])
            # Create HTML fraction using CSS
            result_parts.append(
                f'<span class="fraction">'
                f'<span class="frac-num">{escape(num)}</span>'
                f'<span class="frac-den">{escape(den)}</span>'
                f'</span>'
            )
        return ''.join(result_parts)

    # Fallback: just extract all text
    texts = math_element.findall('.//m:t', namespaces=MATH_NS)
    return escape(''.join([t.text or '' for t in texts]))


def extract_paragraph_html(para_element) -> str:
    """Extract HTML from a paragraph element, including math formulas."""
    html_parts = []

    for child in para_element:
        tag = child.tag.split('}')[-1]  # Get tag name without namespace

        if tag == 'oMath':
            # Math formula
            html_parts.append(extract_math_html(child))
        elif tag == 'r':
            # Regular run - get text
            t_elements = child.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
            text = ''.join([t.text or '' for t in t_elements])
            if text:
                # Check if bold
                rPr = child.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
                is_bold = False
                if rPr is not None:
                    bold_el = rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}b')
                    is_bold = bold_el is not None

                if is_bold:
                    html_parts.append(f'<strong>{escape(text)}</strong>')
                else:
                    html_parts.append(escape(text))
        elif tag == 'oMathPara':
            # Math paragraph - contains oMath elements
            math_elements = child.findall('.//m:oMath', namespaces=MATH_NS)
            for math_el in math_elements:
                html_parts.append(extract_math_html(math_el))

    return ''.join(html_parts)


def extract_cell_html(cell: _Cell) -> str:
    """Convert cell content (paragraphs + nested tables) to HTML."""
    html_parts = []

    # Track which nested tables we've processed
    nested_tables = list(cell.tables)
    table_idx = 0

    # Process cell's XML to maintain order of paragraphs and tables
    for element in cell._tc:
        if element.tag.endswith('}p'):  # Paragraph
            # Extract paragraph content including math
            para_html = extract_paragraph_html(element)
            if para_html.strip():
                html_parts.append(f'<p>{para_html}</p>')
        elif element.tag.endswith('}tbl'):  # Table
            if table_idx < len(nested_tables):
                html_parts.append(extract_nested_table_html(nested_tables[table_idx]))
                table_idx += 1

    return ''.join(html_parts) if html_parts else '&nbsp;'


def generate_html_for_drug(drug_name: str, header_cells: list[str], drug_cells: list[str]) -> str:
    """Generate responsive HTML fragment for a single drug (no full document wrapper)."""
    # Build card sections for each column (skip first column - drug name)
    sections_html = ""
    for i in range(1, len(header_cells)):
        header = escape(header_cells[i])
        content = drug_cells[i] if i < len(drug_cells) else "&nbsp;"
        sections_html += f'''
        <div class="info-section">
            <div class="section-header">{header}</div>
            <div class="section-content">{content}</div>
        </div>'''

    html = f'''<div class="drug-card" data-drug="{escape(drug_name)}">
    <div class="drug-name">{escape(drug_name)}</div>
    {sections_html}
</div>'''
    return html


def generate_standalone_html(drug_name: str, header_cells: list[str], drug_cells: list[str]) -> str:
    """Generate complete standalone HTML document for a single drug."""
    card_html = generate_html_for_drug(drug_name, header_cells, drug_cells)

    html = f'''<!DOCTYPE html>
<html lang="vi">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{escape(drug_name)} - Hướng dẫn hiệu chỉnh liều</title>
    <style>
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}

        body {{
            font-family: system-ui, -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            font-size: 14px;
            line-height: 1.5;
            background: #f5f5f5;
            padding: 16px;
        }}

        .drug-card {{
            background: white;
            border-radius: 8px;
            box-shadow: 0 2px 8px rgba(0,0,0,0.1);
            overflow: hidden;
            max-width: 1200px;
            margin: 0 auto;
        }}

        .drug-name {{
            background: linear-gradient(135deg, #2e7d32, #4caf50);
            color: white;
            font-size: 1.25rem;
            font-weight: 600;
            padding: 16px 20px;
        }}

        .info-section {{
            border-bottom: 1px solid #e0e0e0;
        }}

        .info-section:last-child {{
            border-bottom: none;
        }}

        .section-header {{
            background: #e8f5e9;
            color: #1b5e20;
            font-weight: 600;
            font-size: 0.9rem;
            padding: 12px 16px;
            border-bottom: 1px solid #c8e6c9;
        }}

        .section-content {{
            padding: 16px;
            font-size: 0.95rem;
        }}

        .section-content p {{
            margin: 8px 0;
        }}

        .section-content p:first-child {{
            margin-top: 0;
        }}

        .section-content p:last-child {{
            margin-bottom: 0;
        }}

        /* Math fraction styling */
        .fraction {{
            display: inline-flex;
            flex-direction: column;
            align-items: center;
            vertical-align: middle;
            margin: 0 4px;
            font-size: 0.85em;
        }}

        .frac-num {{
            border-bottom: 1px solid #333;
            padding: 0 4px 2px 4px;
        }}

        .frac-den {{
            padding: 2px 4px 0 4px;
        }}

        /* Nested tables for ClCr dosing */
        .nested-table {{
            width: 100%;
            border-collapse: collapse;
            margin: 12px 0;
            font-size: 0.85rem;
        }}

        .nested-table td {{
            border: 1px solid #ccc;
            padding: 8px 12px;
            text-align: center;
        }}

        .nested-table tr:first-child td {{
            background: #f5f5f5;
            font-weight: 600;
        }}

        .nested-table tr:nth-child(even) td {{
            background: #fafafa;
        }}

        strong {{
            font-weight: 600;
        }}

        /* Responsive adjustments */
        @media (max-width: 768px) {{
            body {{
                padding: 8px;
            }}

            .drug-name {{
                font-size: 1.1rem;
                padding: 12px 16px;
            }}

            .section-header {{
                font-size: 0.85rem;
                padding: 10px 12px;
            }}

            .section-content {{
                padding: 12px;
                font-size: 0.9rem;
            }}

            .nested-table {{
                font-size: 0.8rem;
            }}

            .nested-table td {{
                padding: 6px 8px;
            }}
        }}

        @media (max-width: 480px) {{
            .nested-table {{
                font-size: 0.75rem;
            }}

            .nested-table td {{
                padding: 4px 6px;
            }}
        }}
    </style>
</head>
<body>
    {card_html}
</body>
</html>'''
    return html


def save_html_file(html: str, output_path: Path) -> None:
    """Save HTML content to file."""
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(html)


def main() -> None:
    """Generate responsive HTML files for all drugs."""
    print("Generating drug HTML files from DOCX...")

    # Create output directory
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    # Load DOCX
    doc = Document(DOCX_FILE)
    main_table = doc.tables[0]

    # Extract header row
    header_row = main_table.rows[0]
    header_cells = [cell.text.strip().replace("\n", " ") for cell in header_row.cells]
    print(f"Header: {header_cells}")

    # Count drugs
    drug_count = len(main_table.rows) - 1
    print(f"Found {drug_count} drugs to process")

    success_count = 0
    for row_idx, row in enumerate(main_table.rows[1:], start=1):
        drug_name = row.cells[0].text.strip()
        print(f"\n[{row_idx}/{drug_count}] Processing: {drug_name}...")

        try:
            # Extract cell HTML for each column
            drug_cells = []
            for cell in row.cells:
                cell_html = extract_cell_html(cell)
                drug_cells.append(cell_html)

            # Generate standalone HTML
            html = generate_standalone_html(drug_name, header_cells, drug_cells)

            # Generate filename and save
            safe_name = sanitize_filename(drug_name)
            output_path = OUTPUT_DIR / f"{safe_name}.html"

            save_html_file(html, output_path)
            print(f"  Created: {output_path.name}")
            success_count += 1

        except Exception as e:
            print(f"  Error: {e}")

    print(f"\nDone! Created {success_count}/{drug_count} HTML files in {OUTPUT_DIR}/")


if __name__ == "__main__":
    main()
